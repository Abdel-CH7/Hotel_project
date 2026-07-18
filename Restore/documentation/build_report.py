from __future__ import annotations

import copy
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from PIL import Image


SOURCE_REPORT = Path(
    r"C:\Users\admin\Downloads\Rapport_PFF_Gestion_Hoteliere_TDM_Figures_Active.docx"
)
DOCUMENTATION_DIR = Path(__file__).resolve().parent
DIAGRAMS = DOCUMENTATION_DIR / "diagrams" / "png"
OUTPUT_DIR = DOCUMENTATION_DIR / "output"
OUTPUT_REPORT = OUTPUT_DIR / "Rapport_PFF_Gestion_Hoteliere_Diagrammes_Finalises.docx"


FIGURES = {
    1: ("01-use-cases.png", "Diagramme de cas d’utilisation"),
    2: ("02-architecture.png", "Architecture générale de la solution"),
    3: ("03-domain-classes.png", "Diagramme de classes / domaine"),
    4: ("04-database-erd.png", "Modèle relationnel de la base de données"),
    5: ("05-auth-flow.png", "Flux d’authentification et d’autorisation"),
    6: ("06-reservation-workflow.png", "Workflow principal de création d’une réservation"),
    7: ("07-reservation-sequence.png", "Séquence : création d’une réservation"),
    8: ("08-payment-sequence.png", "Séquence : enregistrement et annulation d’un paiement"),
    9: ("09-complaint-sequence.png", "Séquence : traitement d’une réclamation"),
}


def iter_all_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def find_paragraph(document: Document, exact_text: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def find_table(document: Document, starts_with: str):
    for table in document.tables:
        text = " ".join(cell.text for row in table.rows for cell in row.cells).strip()
        if text.startswith(starts_with):
            return table
    raise ValueError(f"Table not found: {starts_with}")


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_keep(paragraph, *, with_next=False, together=False, page_break_before=False):
    fmt = paragraph.paragraph_format
    fmt.keep_with_next = with_next
    fmt.keep_together = together
    fmt.page_break_before = page_break_before


def add_complex_field(paragraph, instruction: str, placeholder: str = ""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    run._r.append(begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    instr_run._r.append(instr)

    sep_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(separate)

    if placeholder:
        paragraph.add_run(placeholder)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def set_caption(paragraph, title: str):
    clear_paragraph(paragraph)
    try:
        paragraph.style = "Caption"
    except KeyError:
        pass
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Figure ")
    add_complex_field(paragraph, r"SEQ Figure \* ARABIC", "0")
    paragraph.add_run(f" — {title}")
    set_keep(paragraph, together=True)


def set_image_alt_text(run, title: str):
    inline = run._r.xpath(".//wp:inline")
    if not inline:
        return
    doc_pr = inline[0].find(qn("wp:docPr"))
    if doc_pr is not None:
        doc_pr.set("title", title)
        doc_pr.set("descr", title)


def image_size(path: Path, max_width_cm: float, max_height_cm: float):
    with Image.open(path) as image:
        width, height = image.size
    ratio = width / height
    width_cm = max_width_cm
    height_cm = width_cm / ratio
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * ratio
    return Cm(width_cm), Cm(height_cm)


def add_figure_to_paragraph(paragraph, image_path: Path, title: str, *, landscape=False):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    max_width = 25.5 if landscape else 15.8
    max_height = 16.4 if landscape else 20.0
    if image_path.name == "05-auth-flow.png":
        max_height = 13.4
    width, height = image_size(image_path, max_width, max_height)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=width, height=height)
    set_image_alt_text(run, title)
    set_keep(paragraph, with_next=True, together=True)


def set_table_width(table, width_cm: float):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width_twips = int(Cm(width_cm).twips)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_twips))
    for grid_col in table._tbl.tblGrid.gridCol_lst:
        grid_col.set(qn("w:w"), str(width_twips))
    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for cell in row.cells:
            cell.width = Cm(width_cm)
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width_twips))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def replace_table_with_figure(table, figure_number: int, *, landscape=False):
    image_name, title = FIGURES[figure_number]
    cell = table.cell(0, 0)
    cell.text = ""
    figure_paragraph = cell.paragraphs[0]
    add_figure_to_paragraph(figure_paragraph, DIAGRAMS / image_name, title, landscape=landscape)
    caption = cell.add_paragraph()
    set_caption(caption, title)
    set_table_width(table, 25.5 if landscape else 15.8)


def insert_figure_after(element, figure_number: int, *, landscape=False):
    image_name, title = FIGURES[figure_number]
    document = element._parent if hasattr(element, "_parent") else None
    # Temporary paragraphs are created at the end, then moved to the requested location.
    doc = element.part.document
    figure_paragraph = doc.add_paragraph()
    caption = doc.add_paragraph()
    add_figure_to_paragraph(
        figure_paragraph,
        DIAGRAMS / image_name,
        title,
        landscape=landscape,
    )
    set_caption(caption, title)
    element._element.addnext(figure_paragraph._p)
    figure_paragraph._p.addnext(caption._p)


def replace_manual_index_table(table, field_instruction: str):
    document = table.part.document
    paragraph = document.add_paragraph()
    add_complex_field(paragraph, field_instruction, "Mise à jour du champ en cours…")
    table._element.addprevious(paragraph._p)
    table._element.getparent().remove(table._element)


def insert_section_break_before(paragraph, source_sect_pr, *, landscape=False, restart_page_number=False):
    break_paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "1")
    spacing.set(qn("w:lineRule"), "exact")
    p_pr.append(spacing)
    run_properties = OxmlElement("w:rPr")
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), "2")
    run_properties.append(font_size)
    complex_font_size = OxmlElement("w:szCs")
    complex_font_size.set(qn("w:val"), "2")
    run_properties.append(complex_font_size)
    p_pr.append(run_properties)
    sect_pr = copy.deepcopy(source_sect_pr)

    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is not None and not restart_page_number:
        sect_pr.remove(pg_num)
    elif restart_page_number:
        if pg_num is None:
            pg_num = OxmlElement("w:pgNumType")
            sect_pr.append(pg_num)
        pg_num.set(qn("w:start"), "1")

    pg_sz = sect_pr.find(qn("w:pgSz"))
    if pg_sz is None:
        pg_sz = OxmlElement("w:pgSz")
        sect_pr.append(pg_sz)
    if landscape:
        width = pg_sz.get(qn("w:w"), "11906")
        height = pg_sz.get(qn("w:h"), "16838")
        pg_sz.set(qn("w:w"), height)
        pg_sz.set(qn("w:h"), width)
        pg_sz.set(qn("w:orient"), "landscape")
    else:
        pg_sz.attrib.pop(qn("w:orient"), None)

    p_pr.append(sect_pr)
    break_paragraph.append(p_pr)
    paragraph._p.addprevious(break_paragraph)


def remove_empty_paragraph_after(element):
    sibling = element._element.getnext()
    if sibling is None or sibling.tag != qn("w:p"):
        return
    if sibling.xpath(".//w:drawing | .//w:pict | .//w:sectPr"):
        return
    if "".join(sibling.itertext()).strip():
        return
    sibling.getparent().remove(sibling)


def remove_final_page_restart(document: Document):
    body_sect_pr = document.element.body.sectPr
    pg_num = body_sect_pr.find(qn("w:pgNumType"))
    if pg_num is not None:
        body_sect_pr.remove(pg_num)


def convert_screenshot_captions(document: Document):
    screenshot_pattern = re.compile(r"^Figure\s+\d+\s+—\s+(.+?)\s+Capture de la version finale à insérer$")
    converted = 0
    for table in document.tables:
        if len(table.rows) != 1 or len(table.rows[0].cells) != 1:
            continue
        cell = table.cell(0, 0)
        compact = " ".join(cell.text.split())
        match = screenshot_pattern.match(compact)
        if not match:
            continue
        title = match.group(1)
        cell.text = ""
        caption = cell.paragraphs[0]
        set_caption(caption, title)
        placeholder = cell.add_paragraph("Capture de la version finale à insérer")
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_table_width(table, 15.8)
        converted += 1
    if converted != 29:
        raise RuntimeError(f"Expected 29 screenshot captions, converted {converted}")


def add_outline_levels(document: Document):
    chapter_titles = {
        "CONTEXTE, OBJECTIFS ET CAHIER DES CHARGES",
        "ANALYSE ET CONCEPTION",
        "RÉALISATION TECHNIQUE",
        "PRÉSENTATION FONCTIONNELLE DE L’APPLICATION",
        "BILAN, LIMITES ET PERSPECTIVES",
    }
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        level = None
        if text in chapter_titles:
            level = 0
        elif re.match(r"^\d+\.\d+\s", text):
            level = 1
        if level is None:
            continue
        p_pr = paragraph._p.get_or_add_pPr()
        outline = p_pr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            p_pr.append(outline)
        outline.set(qn("w:val"), str(level))


def enable_field_updates(document: Document):
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def main():
    if not SOURCE_REPORT.exists():
        raise FileNotFoundError(SOURCE_REPORT)
    for image_name, _ in FIGURES.values():
        image = DIAGRAMS / image_name
        if not image.exists():
            raise FileNotFoundError(image)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE_REPORT, OUTPUT_REPORT)
    document = Document(OUTPUT_REPORT)

    # Replace manual indexes with live Word fields.
    replace_manual_index_table(
        find_table(document, "Dédicace"),
        r'TOC \o "1-3" \h \z \u',
    )
    replace_manual_index_table(
        find_table(document, "Figure 1 — Diagramme de cas d’utilisation"),
        r'TOC \h \z \c "Figure"',
    )

    # Figure 1 remains in the portrait part of Chapter 2.
    replace_table_with_figure(
        find_table(document, "Figure 1 — Diagramme de cas d’utilisation"),
        1,
        landscape=False,
    )

    # Architecture through the sequence diagrams use one landscape section.
    architecture_heading = find_paragraph(document, "2.4 Architecture générale")
    chapter_three = find_paragraph(document, "CHAPITRE 3")
    source_sect_pr = copy.deepcopy(document.element.body.sectPr)
    insert_section_break_before(
        architecture_heading,
        source_sect_pr,
        landscape=False,
        restart_page_number=True,
    )
    insert_section_break_before(
        chapter_three,
        source_sect_pr,
        landscape=True,
        restart_page_number=False,
    )
    remove_final_page_restart(document)

    # Replace the three existing embedded technical images.
    for paragraph_text, figure_number in [
        ("Figure 2 — Architecture générale de la solution", 2),
        ("Figure 5 — Flux d’authentification et d’autorisation", 5),
        ("Figure 6 — Workflow principal de création d’une réservation", 6),
    ]:
        caption = find_paragraph(document, paragraph_text)
        figure_paragraph = caption._p.getprevious()
        while figure_paragraph is not None and figure_paragraph.tag != qn("w:p"):
            figure_paragraph = figure_paragraph.getprevious()
        if figure_paragraph is None:
            raise RuntimeError(f"No image paragraph before {paragraph_text}")
        from docx.text.paragraph import Paragraph

        figure = Paragraph(figure_paragraph, caption._parent)
        image_name, title = FIGURES[figure_number]
        add_figure_to_paragraph(figure, DIAGRAMS / image_name, title, landscape=True)
        set_caption(caption, title)

    replace_table_with_figure(find_table(document, "Figure 3 —"), 3, landscape=True)
    replace_table_with_figure(find_table(document, "Figure 4 —"), 4, landscape=True)
    figure_seven_table = find_table(document, "Figure 7 —")
    replace_table_with_figure(figure_seven_table, 7, landscape=True)
    remove_empty_paragraph_after(figure_seven_table)
    figure_eight_table = find_table(document, "Figure 8 —")
    replace_table_with_figure(figure_eight_table, 8, landscape=True)
    insert_figure_after(figure_eight_table, 9, landscape=True)
    figure_nine_caption = figure_eight_table._element.getnext().getnext()
    from docx.text.paragraph import Paragraph

    remove_empty_paragraph_after(Paragraph(figure_nine_caption, document._body))

    sequence_heading = find_paragraph(document, "2.8 Diagrammes de séquence à finaliser")
    sequence_heading.text = "2.8 Diagrammes de séquence"
    from docx.text.paragraph import Paragraph

    next_paragraph = sequence_heading._p.getnext()
    while next_paragraph is not None and next_paragraph.tag != qn("w:p"):
        next_paragraph = next_paragraph.getnext()
    if next_paragraph is None:
        raise RuntimeError("Sequence description paragraph not found")
    sequence_description = Paragraph(next_paragraph, sequence_heading._parent)
    sequence_description.text = (
        "Les diagrammes de séquence suivants décrivent les échanges réellement implémentés "
        "entre l’utilisateur, le frontend React, l’API Laravel et la base de données. Ils "
        "couvrent la création d’une réservation, l’enregistrement et l’annulation d’un "
        "paiement, ainsi que le traitement d’une réclamation."
    )

    for paragraph in iter_all_paragraphs(document):
        if "Cette version constitue une base entièrement réorientée" in paragraph.text:
            paragraph.text = (
                "Statut de cette version\n"
                "Cette version constitue la version scolaire finale du projet. Les emplacements "
                "marqués dans le chapitre fonctionnel doivent être remplacés par les captures "
                "d’écran définitives avant la remise."
            )
            break

    # Chapter 4 screenshot placeholders stay in place but become Figures 10–38 automatically.
    convert_screenshot_captions(document)
    add_outline_levels(document)
    enable_field_updates(document)

    document.save(OUTPUT_REPORT)
    print(OUTPUT_REPORT)


if __name__ == "__main__":
    main()
