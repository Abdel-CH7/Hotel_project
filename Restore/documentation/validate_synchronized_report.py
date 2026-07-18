from __future__ import annotations

import re
import zipfile
from pathlib import Path

import fitz
from docx import Document

from synchronize_final_report import drawing_hashes_for_screenshots


ROOT = Path(__file__).resolve().parent
BASELINE = ROOT / "Rapport_PFF_Gestion_Hoteliere_Theme_HMS_A_METTRE_A_JOUR.docx"
DOCX = ROOT / "Rapport_PFF_Gestion_Hoteliere_Theme_HMS_Final_Synchronise.docx"
PDF = ROOT / "Rapport_PFF_Gestion_Hoteliere_Theme_HMS_Final_Synchronise.pdf"

OBSOLETE = (
    "220 tests",
    "1 481 assertions",
    "suppression des réservations",
    "consultation, modification et suppression",
    "DELETE /api/reservations",
    "Annulation effectuée depuis l’ancienne interface",
)

REQUIRED = (
    "253 tests",
    "1 754 assertions",
    "252 routes",
    "Service dégradé",
    "Chambre indisponible",
    "EtatChambre reste la seule autorité",
    "Les réservations existantes ne sont jamais annulées ni modifiées automatiquement",
    "Le passage de En cours à Traité exige l’enregistrement d’une réponse",
    "le passage de Traité à Résolu exige qu’une réponse existe déjà",
)


def all_document_text(document: Document) -> str:
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        chunks.extend(cell.text for row in table.rows for cell in row.cells)
    return "\n".join(chunks)


def centered_footer_numbers(page: fitz.Page) -> list[str]:
    numbers = []
    for word in page.get_text("words"):
        x0, y0, x1, _y1, text, *_rest = word
        centered = abs(((x0 + x1) / 2) - (page.rect.width / 2)) < 55
        if y0 > page.rect.height * 0.9 and centered and re.fullmatch(r"\d+", text):
            numbers.append(text)
    return numbers


def main() -> None:
    for path in (BASELINE, DOCX, PDF):
        if not path.exists() or path.stat().st_size == 0:
            raise FileNotFoundError(path)

    baseline = Document(BASELINE)
    final = Document(DOCX)
    text = all_document_text(final)

    for phrase in OBSOLETE:
        if phrase in text:
            raise AssertionError(f"Obsolete phrase remains: {phrase}")
    for phrase in REQUIRED:
        if phrase.casefold() not in text.casefold():
            raise AssertionError(f"Required phrase missing: {phrase}")

    if drawing_hashes_for_screenshots(baseline) != drawing_hashes_for_screenshots(final):
        raise AssertionError("One or more Chapter 4 screenshot images changed")

    with zipfile.ZipFile(DOCX) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        media_count = len([name for name in archive.namelist() if name.startswith("word/media/")])
    if document_xml.count("SEQ Figure") != 39:
        raise AssertionError("The DOCX does not contain exactly 39 figure sequence fields")
    if document_xml.count("TOC ") != 2:
        raise AssertionError("The DOCX must contain one TOC and one list-of-figures field")
    if document_xml.count("<w:drawing") != 41:
        raise AssertionError("The DOCX must display exactly 41 images")

    pdf = fitz.open(PDF)
    if pdf.page_count != 53:
        raise AssertionError(f"Expected 53 PDF pages, found {pdf.page_count}")
    if len(pdf.get_toc()) == 0:
        raise AssertionError("The PDF has no heading bookmarks")

    body_text = "\n".join(pdf[index].get_text() for index in range(7, pdf.page_count))
    caption_numbers = [int(number) for number in re.findall(r"Figure\s+(\d+)\s+—", body_text)]
    if caption_numbers != list(range(1, 40)):
        raise AssertionError(f"Unexpected figure caption sequence: {caption_numbers}")

    figure_list_page = next(
        page
        for page in pdf
        if "Liste des figures" in page.get_text() and "Figure 1 —" in page.get_text()
    )
    figure_list_numbers = [
        int(number)
        for number in re.findall(r"Figure\s+(\d+)\s+—", figure_list_page.get_text())
    ]
    if figure_list_numbers != list(range(1, 40)):
        raise AssertionError("The list of figures is incomplete or out of sequence")

    figure_links = [link for link in figure_list_page.get_links() if link.get("kind") == fitz.LINK_GOTO]
    if len(figure_links) != 39:
        raise AssertionError(f"Expected 39 figure-list links, found {len(figure_links)}")

    toc_pages = [page for page in pdf if "Table des matières" in page.get_text() or any(
        marker in page.get_text() for marker in ("3.7 Tests et validation", "Conclusion générale")
    )]
    toc_links = sum(
        1
        for page in toc_pages
        for link in page.get_links()
        if link.get("kind") == fitz.LINK_GOTO
    )
    if toc_links == 0:
        raise AssertionError("The table of contents has no internal links")

    if centered_footer_numbers(pdf[0]):
        raise AssertionError("The cover page has a visible page number")
    for page_index in range(1, pdf.page_count):
        expected = str(page_index)
        if expected not in centered_footer_numbers(pdf[page_index]):
            raise AssertionError(f"Visible page number {expected} missing on PDF page {page_index + 1}")

    blank_pages = [
        index + 1
        for index, page in enumerate(pdf)
        if len(page.get_text().strip()) < 30 and not page.get_images(full=True)
    ]
    if blank_pages:
        raise AssertionError(f"Accidental blank pages found: {blank_pages}")

    print({
        "pages": pdf.page_count,
        "figures": len(caption_numbers),
        "toc_links": toc_links,
        "figure_links": len(figure_links),
        "heading_bookmarks": len(pdf.get_toc()),
        "displayed_images": document_xml.count("<w:drawing"),
        "packaged_media": media_count,
        "screenshots_preserved": len(drawing_hashes_for_screenshots(final)),
        "page_numbering": "cover hidden; 1-52 continuous",
    })


if __name__ == "__main__":
    main()
