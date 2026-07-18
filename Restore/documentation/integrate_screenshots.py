from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from PIL import Image


SOURCE_REPORT = Path(
    r"C:\Users\admin\Downloads\Rapport_PFF_Gestion_Hoteliere_Diagrammes_Finalises.docx"
)
ROOT = Path(__file__).resolve().parent
SCREENSHOTS = ROOT / "report" / "screenshots"
OUTPUT_DIR = ROOT / "output"
OUTPUT_REPORT = OUTPUT_DIR / "Rapport_PFF_Gestion_Hoteliere_Captures_Finalisees.docx"

SCREENSHOT_PLACEHOLDER = "Capture de la version finale à insérer"

SCREENSHOT_FILES = {
    10: "10-login.png",
    11: "11-navigation.png",
    12: "12-profile.png",
    13: "13-users-list.png",
    14: "14-user-form.png",
    15: "15-dashboard.png",
    16: "16-clients-particuliers.png",
    17: "17-client-particulier-form.png",
    18: "18-clients-societes.png",
    19: "19-client-societe-form.png",
    20: "20-chambres.png",
    21: "21-types-chambres.png",
    22: "22-etat-chambres.png",
    23: "23-periodes-tarifaires.png",
    24: "24-tarifs-chambres.png",
    25: "25-tarifs-repas.png",
    26: "26-reductions.png",
    27: "27-reservations.png",
    28: "28-reservation-client-dates.png",
    29: "29-reservation-rooms.png",
    30: "30-reservation-services.png",
    31: "31-payment-history.png",
    32: "32-payment-receipt.png",
    33: "33-payment-summary.png",
    34: "34-reclamations.png",
    35: "35-reclamation-workflow.png",
    36: "36-equipements.png",
    37: "37-equipment-form.png",
    38: "38-validation-errors.png",
}


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def caption_title(paragraph) -> str:
    text = " ".join(paragraph.text.split())
    match = re.search(r"(?:—|-)\s*(.+)$", text)
    return match.group(1).strip() if match else text


def set_image_alt_text(run, title: str):
    inline = run._r.xpath(".//wp:inline")
    if not inline:
        return
    doc_pr = inline[0].find(qn("wp:docPr"))
    if doc_pr is not None:
        doc_pr.set("title", title)
        doc_pr.set("descr", title)


def image_size(path: Path, max_width_cm: float = 15.8, max_height_cm: float = 19.0):
    with Image.open(path) as image:
        width_px, height_px = image.size
    ratio = width_px / height_px
    width_cm = max_width_cm
    height_cm = width_cm / ratio
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * ratio
    return Cm(width_cm), Cm(height_cm)


def set_table_width(table, width_cm: float = 15.8):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width_twips = int(Cm(width_cm).twips)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_twips))
    for grid_col in table._tbl.tblGrid.gridCol_lst:
        grid_col.set(qn("w:w"), str(width_twips))
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width_twips))


def enable_field_updates(document: Document):
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def finalize_report_wording(document: Document):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "Annexe B — Captures et diagrammes à fournir":
            for run in paragraph.runs:
                if "Captures et diagrammes à fournir" in run.text:
                    run.text = run.text.replace(
                        "Captures et diagrammes à fournir",
                        "Captures et diagrammes intégrés",
                    )

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "Statut de cette version" not in paragraph.text:
                        continue
                    for run in paragraph.runs:
                        if "Les emplacements marqués" in run.text:
                            run.text = (
                                "\nCette version constitue la version scolaire finale du projet. "
                                "Les diagrammes techniques et les captures d’écran définitives "
                                "ont été intégrés et vérifiés avant la remise."
                            )


def insert_screenshot(table, figure_number: int, image_path: Path):
    cell = table.cell(0, 0)
    caption = next(
        (
            paragraph
            for paragraph in cell.paragraphs
            if "SEQ Figure" in paragraph._p.xml or paragraph.text.strip().startswith("Figure ")
        ),
        None,
    )
    placeholder = next(
        (paragraph for paragraph in cell.paragraphs if SCREENSHOT_PLACEHOLDER in paragraph.text),
        None,
    )
    if caption is None or placeholder is None:
        raise RuntimeError(f"Figure {figure_number}: caption or placeholder not found")

    title = caption_title(caption)
    clear_paragraph(placeholder)
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width, height = image_size(image_path)
    run = placeholder.add_run()
    run.add_picture(str(image_path), width=width, height=height)
    set_image_alt_text(run, title)
    placeholder.paragraph_format.keep_with_next = True
    placeholder.paragraph_format.keep_together = True
    caption.paragraph_format.keep_together = True

    # The baseline stores the caption before the placeholder. Move the image
    # paragraph in front without recreating the live SEQ field or bookmarks.
    caption._p.addprevious(placeholder._p)
    set_table_width(table)


def main():
    if not SOURCE_REPORT.exists():
        raise FileNotFoundError(SOURCE_REPORT)
    missing = [str(SCREENSHOTS / name) for name in SCREENSHOT_FILES.values() if not (SCREENSHOTS / name).exists()]
    if missing:
        raise FileNotFoundError("Missing screenshots:\n" + "\n".join(missing))

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document(SOURCE_REPORT)
    placeholder_tables = [
        table
        for table in document.tables
        if SCREENSHOT_PLACEHOLDER in " ".join(cell.text for row in table.rows for cell in row.cells)
    ]
    if len(placeholder_tables) != 29:
        raise RuntimeError(f"Expected 29 screenshot placeholders, found {len(placeholder_tables)}")

    for figure_number, table in zip(range(10, 39), placeholder_tables):
        insert_screenshot(table, figure_number, SCREENSHOTS / SCREENSHOT_FILES[figure_number])

    finalize_report_wording(document)
    enable_field_updates(document)
    document.save(OUTPUT_REPORT)
    print(OUTPUT_REPORT)


if __name__ == "__main__":
    main()
