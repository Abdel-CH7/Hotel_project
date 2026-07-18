from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.text.paragraph import Paragraph
from PIL import Image


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "Rapport_PFF_Gestion_Hoteliere_Theme_HMS_A_METTRE_A_JOUR.docx"
OUTPUT = ROOT / "Rapport_PFF_Gestion_Hoteliere_Theme_HMS_Final_Synchronise.docx"
DIAGRAMS = ROOT / "diagrams" / "png"

UPDATED_FIGURES = {
    "Diagramme de cas d’utilisation": "01-use-cases.png",
    "Diagramme de classes / domaine": "03-domain-classes.png",
    "Modèle relationnel de la base de données": "04-database-erd.png",
    "Workflow principal de création d’une réservation": "06-reservation-workflow.png",
    "Séquence : création d’une réservation": "07-reservation-sequence.png",
}

SCREENSHOT_TITLES = {
    "Page de connexion",
    "Navigation et identité de l’utilisateur",
    "Page Mon profil",
    "Liste et filtres des utilisateurs",
    "Création ou modification d’un utilisateur",
    "Tableau de bord et statistiques",
    "Liste des clients particuliers",
    "Formulaire client particulier et validation",
    "Liste des clients sociétés",
    "Formulaire client société et contacts",
    "Gestion des chambres",
    "Gestion des types de chambres",
    "État des chambres",
    "Périodes tarifaires",
    "Tarifs des chambres",
    "Tarifs des repas",
    "Plans et types de réductions",
    "Liste des réservations",
    "Sélection du client et des dates",
    "Affectation multi-chambres et occupation",
    "Repas, politique de paiement et calcul",
    "Historique des paiements et audit",
    "Reçu de paiement",
    "Récapitulatif de règlement",
    "Liste et filtres des réclamations",
    "Détail et workflow de traitement",
    "Liste des équipements",
    "Formulaire d’équipement et localisation conditionnelle",
    "Exemple de validation inline et champs obligatoires",
}


def iter_all_paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def normalized(text: str) -> str:
    return " ".join(text.split())


def paragraph_exact(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if normalized(paragraph.text) == normalized(text):
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def table_starting(document: Document, text: str):
    for table in document.tables:
        compact = normalized(" ".join(cell.text for row in table.rows for cell in row.cells))
        if compact.startswith(normalized(text)):
            return table
    raise ValueError(f"Table not found: {text}")


def set_paragraph(paragraph: Paragraph, text: str) -> Paragraph:
    paragraph.text = text
    return paragraph


def insert_after(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.style = style or paragraph.style
    inserted.add_run(text)
    return inserted


def set_cell(cell, text: str) -> None:
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def append_table_row(table, values: tuple[str, ...]) -> None:
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)


def keep_table_rows_together(table) -> None:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_complex_field(paragraph: Paragraph, instruction: str, placeholder: str = "0") -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = f" {instruction} "
    instruction_run._r.append(instruction_text)

    separator_run = paragraph.add_run()
    separator = OxmlElement("w:fldChar")
    separator.set(qn("w:fldCharType"), "separate")
    separator_run._r.append(separator)
    paragraph.add_run(placeholder)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def set_caption(paragraph: Paragraph, title: str) -> None:
    clear_paragraph(paragraph)
    paragraph.style = "Caption"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Figure ")
    add_complex_field(paragraph, r"SEQ Figure \* ARABIC")
    paragraph.add_run(f" — {title}")
    paragraph.paragraph_format.keep_together = True


def image_dimensions(path: Path, max_width_cm: float, max_height_cm: float):
    with Image.open(path) as image:
        width_px, height_px = image.size
    ratio = width_px / height_px
    width_cm = max_width_cm
    height_cm = width_cm / ratio
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * ratio
    return Cm(width_cm), Cm(height_cm)


def set_image_alt(run, title: str) -> None:
    inline = run._r.xpath(".//wp:inline")
    if not inline:
        return
    doc_pr = inline[0].find(qn("wp:docPr"))
    if doc_pr is not None:
        doc_pr.set("title", title)
        doc_pr.set("descr", title)


def caption_by_title(document: Document, title: str) -> Paragraph:
    for paragraph in iter_all_paragraphs(document):
        if title in normalized(paragraph.text) and "SEQ Figure" in paragraph._p.xml:
            return paragraph
    raise ValueError(f"Figure caption not found: {title}")


def image_paragraph_before(caption: Paragraph) -> Paragraph:
    sibling = caption._p.getprevious()
    while sibling is not None:
        if sibling.tag == qn("w:p") and sibling.xpath(".//w:drawing | .//w:pict"):
            return Paragraph(sibling, caption._parent)
        sibling = sibling.getprevious()
    raise RuntimeError(f"No image paragraph before caption: {caption.text}")


def replace_figure_image(document: Document, title: str, image_path: Path) -> None:
    caption = caption_by_title(document, title)
    paragraph = image_paragraph_before(caption)
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    max_width, max_height = (15.8, 19.5) if title == "Diagramme de cas d’utilisation" else (25.5, 16.4)
    width, height = image_dimensions(image_path, max_width, max_height)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=width, height=height)
    set_image_alt(run, title)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True


def set_table_width(table, width_cm: float) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width_twips = int(Cm(width_cm).twips)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_twips))
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width_twips))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def insert_equipment_impact_figure(document: Document) -> None:
    title = "Séquence : impact d’un équipement sur une chambre"
    image_path = DIAGRAMS / "10-equipment-impact-sequence.png"
    complaint_caption = caption_by_title(document, "Séquence : traitement d’une réclamation")

    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    image_paragraph = cell.paragraphs[0]
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width, height = image_dimensions(image_path, 25.5, 16.4)
    run = image_paragraph.add_run()
    run.add_picture(str(image_path), width=width, height=height)
    set_image_alt(run, title)
    image_paragraph.paragraph_format.keep_with_next = True
    image_paragraph.paragraph_format.keep_together = True
    caption = cell.add_paragraph()
    set_caption(caption, title)
    set_table_width(table, 25.5)

    complaint_caption._p.addnext(table._tbl)


def drawing_hashes_for_screenshots(document: Document) -> dict[str, str]:
    hashes: dict[str, str] = {}
    for table in document.tables:
        if "SEQ Figure" not in table._tbl.xml:
            continue
        text = normalized(" ".join(cell.text for row in table.rows for cell in row.cells))
        title = next((candidate for candidate in SCREENSHOT_TITLES if candidate in text), None)
        if title is None:
            continue
        blips = table._tbl.xpath(".//a:blip")
        if len(blips) != 1:
            raise RuntimeError(f"Expected one screenshot image for {title}, found {len(blips)}")
        relation_id = blips[0].get(qn("r:embed"))
        part = table.part.related_parts[relation_id]
        hashes[title] = hashlib.sha256(part.blob).hexdigest()
    if set(hashes) != SCREENSHOT_TITLES:
        missing = sorted(SCREENSHOT_TITLES - set(hashes))
        raise RuntimeError(f"Expected 29 screenshot images; missing: {missing}")
    return hashes


def enable_field_updates(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def update_report_text(document: Document) -> None:
    set_paragraph(
        paragraph_exact(
            document,
            "Le système calcule la disponibilité réelle des chambres en tenant compte des réservations qui se chevauchent et des maintenances actives. Il prend en charge les réservations multi-chambres, l’occupation, les repas, les périodes tarifaires, les réductions et les politiques de paiement. Les paiements sont tracés par utilisateur, peuvent être annulés sans effacer l’historique et donnent lieu à des reçus et récapitulatifs imprimables ou exportables en PDF.",
        ),
        "Le système calcule la disponibilité réelle des chambres en tenant compte des réservations qui se chevauchent et des maintenances actives. L’état d’une chambre combine la propreté, la maintenance et une occupation dérivée des réservations. L’impact d’un équipement distingue le service dégradé de la chambre indisponible, tandis qu’une navigation contextuelle relie chambres, réservations, réclamations et équipements. Les réservations multi-chambres, les repas, les tarifs, les réductions, les politiques de paiement et la traçabilité des règlements sont également pris en charge.",
    )
    set_paragraph(
        paragraph_exact(
            document,
            "The solution uses a React and Vite single-page application connected to a Laravel 11 REST API and a MySQL database. Laravel Sanctum provides token-based authentication, while administrator and staff roles control access to the system. The application also supports multi-room reservations, real-time room availability rules, pricing periods, meals, discounts, payment policies, audit-friendly payment records and printable PDF documents.",
        ),
        "The solution uses a React and Vite single-page application connected to a Laravel 11 REST API and a MySQL database. Laravel Sanctum provides token-based authentication, while administrator and staff roles control access. Room state combines cleanliness, maintenance and reservation-derived occupancy. Equipment can cause a non-blocking degraded service or make a room unavailable through room maintenance, and contextual navigation connects rooms, reservations, complaints and equipment. Multi-room stays, pricing, meals, discounts, payment policies and auditable payment documents are also supported.",
    )

    requirements = table_starting(document, "Module Fonctions principales")
    requirement_updates = {
        "Chambres": "Gestion des chambres physiques, types, capacités, étages et vues ; suivi de la propreté, de la maintenance, de l’occupation actuelle et des équipements affectés.",
        "Réservations": "Sélection du client et des dates, disponibilité, affectation multi-chambres, occupation, avertissements de service dégradé, repas, calcul, politique de paiement, consultation, modification et annulation avec motif.",
        "Réclamations": "Création libre ou depuis une réservation, client dérivé, chambre limitée au séjour, priorité, département, traitement, réponse, résolution, annulation et navigation vers les éléments associés.",
        "Équipements": "Catégories, chambre ou emplacement interne, statut, impact sur la chambre, maintenance associée, acquisition, garantie, fournisseur, prix, document et filtrage par chambre.",
    }
    for row in requirements.rows[1:]:
        key = normalized(row.cells[0].text)
        if key in requirement_updates:
            set_cell(row.cells[1], requirement_updates[key])

    rules = table_starting(document, "Règle métier Comportement attendu")
    for rule in (
        ("Occupation d’une chambre", "L’occupation actuelle est calculée à partir des réservations actives et n’est jamais modifiée manuellement."),
        ("Date de départ", "La date de départ est exclusive : un séjour se terminant aujourd’hui ne bloque pas la chambre aujourd’hui."),
        ("Impact d’un équipement", "Un équipement problématique doit avoir l’impact Aucun, Service dégradé ou Chambre indisponible."),
        ("Autorité de blocage", "Seule la maintenance enregistrée dans l’état de la chambre bloque les nouvelles réservations."),
        ("Réservations existantes", "La mise en maintenance d’une chambre ne supprime, n’annule et ne modifie jamais automatiquement les réservations existantes."),
        ("Annulation de réservation", "Une réservation est annulée par une transition de statut avec un motif obligatoire ; elle n’est pas supprimée."),
        ("Affectation d’un équipement", "Un équipement appartient soit à une chambre réelle, soit à un emplacement interne, jamais aux deux."),
    ):
        append_table_row(rules, rule)

    backend_intro = paragraph_exact(
        document,
        "Le backend suit les conventions Laravel. Les contrôleurs reçoivent les requêtes, les Form Requests valident les données, les modèles Eloquent représentent les entités et les services isolent certaines règles complexes. Les migrations assurent l’évolution de la base de données.",
    )
    insert_after(
        backend_intro,
        "Les règles critiques sont centralisées dans des services dédiés. ReservationAvailabilityService contrôle les chevauchements de réservations et les maintenances actives. EquipmentRoomImpactService normalise l’impact des équipements, verrouille les enregistrements concernés et synchronise les maintenances de chambre dans une transaction. ReclamationService contrôle les transitions de statut et les relations avec la réservation, le client et la chambre. Des verrouillages lockForUpdate et des transactions évitent les états incohérents lors des opérations concurrentes.",
    )
    set_paragraph(
        paragraph_exact(document, "Transactions : utilisées pour les opérations qui doivent rester atomiques, comme les changements touchant le dernier administrateur."),
        "Transactions : utilisées pour les opérations qui doivent rester atomiques, notamment la création et la modification d’une réservation, la synchronisation Chambre/EtatChambre, l’impact bloquant d’un équipement et les changements touchant le dernier administrateur.",
    )

    validation = paragraph_exact(
        document,
        "La qualité du backend a été vérifiée par des tests fonctionnels couvrant l’authentification, les autorisations, les comptes utilisateurs, les réservations, les paiements et les autres modules. Au dernier contrôle global, la suite backend comportait 220 tests et 1 481 assertions réussies. Le frontend a également passé une compilation de production avec succès.",
    )
    set_paragraph(
        validation,
        "La qualité du backend a été vérifiée par des tests fonctionnels couvrant l’authentification, les autorisations, les comptes utilisateurs, les réservations, les paiements et les autres modules. Au dernier contrôle global, la suite backend comportait 253 tests et 1 754 assertions réussies.",
    )
    validation_routes = insert_after(
        validation,
        "L’inventaire comportait 252 routes, sans aucune route DELETE permettant de supprimer ou d’annuler une réservation.",
    )
    insert_after(
        validation_routes,
        "Le frontend a passé une compilation de production avec succès. Le contrôle git diff --check n’a signalé aucune erreur, uniquement des avertissements de conversion de fins de ligne.",
    )

    state = paragraph_exact(
        document,
        "La page État des chambres suit la propreté et la maintenance. La propreté utilise les valeurs Nettoyée et Non nettoyée. Une maintenance active rend la chambre indisponible pour les nouvelles réservations sur la période concernée.",
    )
    set_paragraph(state, "La page État des chambres centralise trois dimensions indépendantes : la propreté, la maintenance et l’occupation commerciale.")
    state = insert_after(state, "La propreté utilise les valeurs Nettoyée et Non nettoyée et conserve, lorsqu’elles sont disponibles, la date du dernier nettoyage ainsi que l’employé concerné.")
    state = insert_after(state, "L’occupation n’est pas saisie manuellement. Elle est calculée à partir des réservations En attente ou Confirmées dont la période couvre la date actuelle. La date de départ est une limite exclusive : une chambre dont le séjour se termine aujourd’hui est considérée comme libre aujourd’hui.")
    state = insert_after(state, "Une maintenance active reste la seule autorité opérationnelle capable de rendre la chambre indisponible pour les nouvelles réservations sur la période concernée.")
    state = insert_after(state, "La page présente aussi un résumé des équipements affectés à chaque chambre et permet d’accéder à leur liste filtrée.")
    insert_after(state, "Des liens contextuels donnent accès à la chambre physique, à ses équipements et à la réservation actuelle lorsqu’elle existe.")

    set_paragraph(
        paragraph_exact(document, "La réservation constitue le cœur du système. L’utilisateur sélectionne d’abord le type de client et le client, puis saisit les dates de séjour. L’API retourne uniquement les chambres réellement disponibles, après exclusion des conflits de réservation et des maintenances actives."),
        "La réservation constitue le cœur du système. L’utilisateur sélectionne le type de client, le client et les dates du séjour. L’API retourne les chambres disponibles après exclusion des chevauchements et des maintenances actives.",
    )
    reservation_end = set_paragraph(
        paragraph_exact(document, "Les repas, les réductions et la politique de paiement participent au calcul. Le système prend en charge les statuts En attente, Confirmé et Annulé, ainsi que la consultation, la modification et la suppression selon les règles existantes."),
        "Les repas, les réductions et la politique de paiement participent au calcul. Le système prend en charge les statuts En attente, Confirmé et Annulé, ainsi que la consultation, la modification et l’annulation contrôlée.",
    )
    reservation_end = insert_after(reservation_end, "Une réservation n’est jamais supprimée depuis l’interface. L’annulation est réalisée par une transition de statut et exige un motif non vide afin de préserver l’historique commercial et les paiements associés.")
    reservation_end = insert_after(reservation_end, "Lorsqu’un équipement affecté à une chambre provoque uniquement un service dégradé, la chambre reste sélectionnable mais un avertissement non bloquant est affiché à l’utilisateur.")
    reservation_end = insert_after(reservation_end, "Lorsqu’un équipement a entraîné une maintenance de chambre active, la chambre est exclue des disponibilités sur la période concernée.")
    insert_after(reservation_end, "Depuis le détail d’une réservation, l’utilisateur peut ouvrir les réclamations associées ou créer une nouvelle réclamation avec la réservation déjà sélectionnée.")

    complaint = paragraph_exact(
        document,
        "Le module Réclamations permet d’enregistrer la demande d’un client, d’attribuer une priorité et un département, puis de suivre le traitement. Les statuts utilisés sont En attente, En cours, Traité, Résolu et Annulé. La résolution exige qu’une réponse existe déjà, ce qui empêche de clôturer un dossier sans trace du traitement effectué.",
    )
    set_paragraph(complaint, "Une réclamation peut être associée à une réservation. Dans ce cas, le client est déterminé automatiquement à partir de la réservation et le choix de la chambre est limité aux chambres appartenant au séjour.")
    complaint = insert_after(complaint, "Une réclamation peut aussi être enregistrée avec un client direct lorsqu’aucune réservation n’est disponible. En revanche, une chambre ne peut pas être associée seule sans contexte de réservation.")
    complaint = insert_after(complaint, "Le workflow utilise les statuts En attente, En cours, Traité, Résolu et Annulé.")
    complaint = insert_after(complaint, "Le passage de En cours à Traité exige l’enregistrement d’une réponse, et le passage de Traité à Résolu exige qu’une réponse existe déjà. Cette règle empêche la clôture d’un dossier sans trace du traitement réalisé.")
    insert_after(complaint, "Les actions contextuelles permettent d’accéder à la réservation, à la chambre physique et à son état lorsque ces relations existent.")

    equipment = paragraph_exact(
        document,
        "La gestion des équipements couvre le nom, le numéro de série, la marque, le modèle, la catégorie, le statut, la date d’acquisition, la garantie, le fournisseur, le prix d’achat et un document éventuel. Un équipement peut être affecté soit à une chambre, soit à un emplacement interne.",
    )
    set_paragraph(equipment, "La gestion des équipements couvre le nom, le numéro de série, la marque, le modèle, la catégorie, le statut, la date d’acquisition, la garantie, le fournisseur, le prix d’achat, les notes et un document éventuel. Un équipement doit être affecté soit à une chambre réelle, soit à un emplacement interne de l’hôtel, mais jamais aux deux simultanément.")
    equipment = set_paragraph(
        paragraph_exact(document, "Les statuts sont présentés sous les libellés En service, En maintenance et Hors service. La garantie est mise en évidence selon qu’elle est valide, proche de l’expiration, expirée ou non renseignée."),
        "L’utilisation d’un emplacement artificiel portant le nom d’une chambre est refusée afin de conserver une relation fiable avec la chambre physique. Les statuts sont présentés sous les libellés En service, En maintenance et Hors service.",
    )
    equipment = insert_after(equipment, "Lorsqu’un équipement de chambre devient problématique, l’utilisateur doit préciser son impact :")
    equipment = insert_after(equipment, "Aucun impact : l’équipement n’affecte pas l’exploitation de la chambre.", "List Bullet")
    equipment = insert_after(equipment, "Service dégradé : la chambre reste réservable, mais un avertissement est présenté lors de son affectation à une réservation.", "List Bullet")
    equipment = insert_after(equipment, "Chambre indisponible : une maintenance de chambre est créée ou réutilisée dans une transaction, et la chambre est bloquée pour les nouvelles réservations pendant la période concernée.", "List Bullet")
    equipment = insert_after(equipment, "Les réservations existantes ne sont jamais annulées ni modifiées automatiquement lorsqu’un équipement devient bloquant. Si une maintenance les chevauche, une confirmation explicite est demandée à l’utilisateur.")
    equipment = insert_after(equipment, "Le retour de l’équipement au statut En service ne remet pas automatiquement la chambre en vente. Son état doit être vérifié manuellement, car EtatChambre reste la seule autorité capable de bloquer ou de libérer une chambre.")
    insert_after(equipment, "La page propose un filtre par chambre, combinable avec le statut, la catégorie et la recherche. Le nombre d’équipements affectés à chaque chambre est affiché dans le sélecteur.")

    search = paragraph_exact(
        document,
        "Les principales listes proposent une recherche, des filtres, une pagination et, lorsque cela est pertinent, des exports Excel/PDF et une impression. Les termes recherchés sont mis en évidence dans les tableaux. Les formulaires utilisent le même style d’erreur dans les différents modules, ce qui réduit les incohérences d’interface.",
    )
    search = insert_after(search, "Une navigation contextuelle relie les chambres, leur état, leurs équipements, les réservations et les réclamations.")
    search = insert_after(search, "Les pages utilisent des paramètres d’URL afin de filtrer une liste ou d’ouvrir une ressource précise sans créer de nouvelles routes frontend.")
    search = insert_after(search, "La fermeture d’un détail ou l’effacement d’un contexte retire uniquement le paramètre concerné et conserve les autres filtres éventuels.")
    insert_after(search, "Ce fonctionnement permet aussi de restaurer le contexte avec les actions Précédent et Suivant du navigateur.")

    result_updates = {
        "Gestion structurée des clients, chambres, tarifs, réservations et services.": "Gestion structurée des clients, chambres, tarifs, réservations et services.",
        "Disponibilité calculée en fonction des réservations et de la maintenance.": "État des chambres enrichi par la propreté, la maintenance, l’occupation actuelle et le résumé des équipements.",
        "Paiements traçables avec documents imprimables et PDF.": "Paiements traçables avec documents imprimables et PDF.",
        "Workflow de réclamation avec règle de résolution.": "Workflow de réclamation avec réponse obligatoire avant traitement final et résolution.",
        "Gestion des équipements et de leur garantie.": "Gestion de l’impact des équipements avec distinction entre service dégradé et chambre indisponible.",
        "Tests backend réussis et build frontend validé.": "Navigation contextuelle entre chambres, états, équipements, réservations et réclamations.",
    }
    final_result = None
    for source, replacement in result_updates.items():
        final_result = set_paragraph(paragraph_exact(document, source), replacement)
    final_result = insert_after(final_result, "Annulation des réservations sans suppression physique et avec conservation obligatoire du motif.", "List Bullet")
    insert_after(final_result, "253 tests backend et 1 754 assertions réussies, compilation frontend validée et aucun endpoint DELETE pour les réservations.", "List Bullet")

    difficulties = table_starting(document, "Difficulté Solution retenue")
    append_table_row(difficulties, (
        "Relations entre modules",
        "Les chambres, leur état, les réservations, les réclamations et les équipements avaient initialement des comportements isolés. Les relations ont été centralisées côté backend et une navigation contextuelle fondée sur des identifiants exacts a été ajoutée.",
    ))
    append_table_row(difficulties, (
        "Compatibilité avec les anciennes données",
        "Certaines données de l’ancienne interface étaient incomplètes pour les nouveaux workflows. Les protections de lecture utiles ont été conservées, les interfaces obsolètes comme DELETE réservation ont été supprimées et un jeu de données final cohérent a été préparé.",
    ))
    keep_table_rows_together(difficulties)

    set_paragraph(
        paragraph_exact(document, "Le travail réalisé a dépassé un simple ensemble de pages CRUD. Il a nécessité l’application de règles métier liées à la disponibilité des chambres, à la tarification, à l’occupation, aux paiements, aux réclamations, aux comptes utilisateurs et à la traçabilité. La séparation entre React/Vite et l’API Laravel 11 a favorisé une architecture claire et maintenable."),
        "Le travail réalisé a dépassé un simple ensemble de pages CRUD. La solution finale assure la cohérence entre l’état des chambres, les réservations, l’impact des équipements, les réclamations et leur navigation contextuelle. La séparation entre React/Vite et l’API Laravel 11 a favorisé une architecture claire et maintenable.",
    )
    set_paragraph(
        paragraph_exact(document, "La mise en place de Laravel Sanctum, des rôles, du statut actif, de la révocation des jetons et des protections du dernier administrateur a renforcé la sécurité. Les composants communs, la validation inline, les exports et les documents de paiement ont amélioré l’expérience utilisateur. Enfin, les tests backend et la compilation de production ont fourni une base de confiance pour la stabilité de la solution."),
        "Le workflow des équipements distingue un service dégradé d’une chambre réellement indisponible sans modifier automatiquement les réservations existantes. Laravel Sanctum, les rôles, le statut actif, les composants communs, la validation inline et les contrôles automatisés fournissent une base fiable pour la stabilité et l’évolution de la solution.",
    )

    endpoints = table_starting(document, "Domaine Endpoints / opérations")
    for row in endpoints.rows[1:]:
        domain = normalized(row.cells[0].text)
        if domain == "Réservations":
            set_cell(row.cells[1], "GET/POST /api/reservations ; GET/PUT /api/reservations/{reservation} ; PATCH /api/reservations/{reservation}/status ; GET /api/reservations/available-rooms ; POST /api/reservations/calculate-price ; paiements associés. Aucune route DELETE n’est exposée ; l’annulation utilise PATCH et exige un motif.")
        elif domain == "Modules métier":
            set_cell(row.cells[0], "Autres modules")
            set_cell(row.cells[1], "Clients, tarifs et statistiques selon les routes protégées existantes.")
    append_table_row(endpoints, ("État des chambres", "GET/POST /api/etat-chambre ; GET/PUT /api/etat-chambre/{num_chambre} ; consultation, propreté, nettoyage, maintenance et occupation actuelle."))
    append_table_row(endpoints, ("Réclamations", "GET/POST /api/reclamations ; GET/PUT /api/reclamations/{reclamation} ; PATCH status/cancel ; contexte de réservation."))
    append_table_row(endpoints, ("Équipements", "GET/POST /api/equipements ; GET/PUT/PATCH/DELETE /api/equipements/{equipement} ; catégories, emplacements, impact sur la chambre et synthèse par chambre."))

    set_paragraph(
        paragraph_exact(document, "Les diagrammes de séquence suivants décrivent les échanges réellement implémentés entre l’utilisateur, le frontend React, l’API Laravel et la base de données. Ils couvrent la création d’une réservation, l’enregistrement et l’annulation d’un paiement, ainsi que le traitement d’une réclamation."),
        "Les diagrammes de séquence suivants décrivent les échanges réellement implémentés entre l’utilisateur, le frontend React, l’API Laravel et la base de données. Ils couvrent la création d’une réservation, l’enregistrement et l’annulation d’un paiement, le traitement d’une réclamation et l’impact contrôlé d’un équipement sur une chambre.",
    )
    set_paragraph(
        paragraph_exact(document, "Diagrammes de séquence pour réservation, paiement et réclamation."),
        "Diagrammes de séquence pour réservation, paiement, réclamation et impact d’un équipement sur une chambre.",
    )
    set_paragraph(
        paragraph_exact(document, "Résultat final de la commande de tests backend et du build frontend."),
        "Résultats finaux : 253 tests backend, 1 754 assertions, 252 routes et build frontend validé.",
    )


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    for image_name in UPDATED_FIGURES.values():
        if not (DIAGRAMS / image_name).exists():
            raise FileNotFoundError(DIAGRAMS / image_name)
    if not (DIAGRAMS / "10-equipment-impact-sequence.png").exists():
        raise FileNotFoundError(DIAGRAMS / "10-equipment-impact-sequence.png")

    document = Document(SOURCE)
    screenshot_hashes = drawing_hashes_for_screenshots(document)

    update_report_text(document)
    for title, image_name in UPDATED_FIGURES.items():
        replace_figure_image(document, title, DIAGRAMS / image_name)
    insert_equipment_impact_figure(document)
    enable_field_updates(document)

    if drawing_hashes_for_screenshots(document) != screenshot_hashes:
        raise RuntimeError("A screenshot image changed during report synchronization")
    seq_count = sum("SEQ Figure" in paragraph._p.xml for paragraph in iter_all_paragraphs(document))
    if seq_count != 39:
        raise RuntimeError(f"Expected 39 figure sequence fields, found {seq_count}")

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
